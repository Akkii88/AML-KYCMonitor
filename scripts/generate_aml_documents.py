import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import PieChart, BarChart, DoughnutChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
import os
from datetime import datetime

# AML/KYC CONTEXT:
# Professional AML management dashboards provide executives with clear visibility into compliance metrics.
# Real banking dashboards use dense, structured layouts with muted corporate colors.
# This dashboard resembles internal MIS reports used for compliance review meetings.

ALERTS_PATH = 'data/cleaned/aml_alerts.csv'
INVESTIGATIONS_PATH = 'data/cleaned/investigation_cases.csv'
CUSTOMERS_PATH = 'data/cleaned/customers_cleaned.csv'
TRANSACTIONS_PATH = 'data/cleaned/transactions_cleaned.csv'

def generate_excel_workbook():
    alerts = pd.read_csv(ALERTS_PATH)
    investigations = pd.read_csv(INVESTIGATIONS_PATH)
    customers = pd.read_csv(CUSTOMERS_PATH)
    transactions = pd.read_csv(TRANSACTIONS_PATH)
    
    os.makedirs('reports', exist_ok=True)
    
    wb = Workbook()
    
    # ========== EXECUTIVE DASHBOARD ==========
    ws = wb.active
    ws.title = "Executive Dashboard"
    
    # Hide gridlines
    ws.sheet_view.showGridLines = False
    
    # Column widths - compact professional layout
    ws.column_dimensions['A'].width = 2
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 16
    ws.column_dimensions['D'].width = 2
    ws.column_dimensions['E'].width = 22
    ws.column_dimensions['F'].width = 16
    ws.column_dimensions['G'].width = 2
    
    # Professional color palette
    NAVY = "1B4F72"
    MUTED_BLUE = "2E86AB"
    SOFT_RED = "C0392B"
    GREY = "566573"
    LIGHT_GREY = "F4F6F7"
    WHITE = "FFFFFF"
    
    header_fill = PatternFill(start_color=NAVY, end_color=NAVY, fill_type="solid")
    header_font = Font(color=WHITE, size=11, bold=True)
    kpi_fill = PatternFill(start_color=LIGHT_GREY, end_color=LIGHT_GREY, fill_type="solid")
    kpi_value_font = Font(size=18, bold=True, color=NAVY)
    kpi_label_font = Font(size=9, color=GREY)
    section_font = Font(size=10, bold=True, color=NAVY)
    thin_border = Border(
        left=Side(style='thin', color='D5D8DC'),
        right=Side(style='thin', color='D5D8DC'),
        top=Side(style='thin', color='D5D8DC'),
        bottom=Side(style='thin', color='D5D8DC')
    )
    
    # ========== TITLE ==========
    ws.merge_cells('B2:F2')
    ws['B2'] = "AML COMPLIANCE MONITORING DASHBOARD"
    ws['B2'].font = Font(size=14, bold=True, color=NAVY)
    ws['B2'].alignment = Alignment(horizontal='center', vertical='center')
    
    ws.merge_cells('B3:F3')
    ws['B3'] = f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ws['B3'].font = Font(size=8, color=GREY)
    ws['B3'].alignment = Alignment(horizontal='center')
    
    # ========== KPI SECTION ==========
    ws['B5'] = "KEY PERFORMANCE INDICATORS"
    ws['B5'].font = header_font
    ws['B5'].fill = header_fill
    ws.merge_cells('B5:F5')
    ws['B5'].alignment = Alignment(horizontal='center', vertical='center')
    
    # KPI values
    kpis = [
        ('Total Alerts', len(alerts)),
        ('Escalated Cases', len(investigations[investigations['escalation_required'] == 'Yes'])),
        ('High-Risk Customers', len(customers[customers['risk_category'] == 'High'])),
        ('Pending KYC', len(customers[customers['kyc_status'] == 'Pending'])),
        ('Fraudulent Transactions', int(transactions['isFraud'].sum()))
    ]
    
    for i, (label, value) in enumerate(kpis):
        col = 2 + i
        cell_label = ws.cell(row=6, column=col, value=label)
        cell_label.font = kpi_label_font
        cell_label.alignment = Alignment(horizontal='center', vertical='center')
        cell_label.fill = kpi_fill
        cell_label.border = thin_border
        
        cell_value = ws.cell(row=7, column=col, value=value)
        cell_value.font = kpi_value_font
        cell_value.alignment = Alignment(horizontal='center', vertical='center')
        cell_value.fill = kpi_fill
        cell_value.border = thin_border
    
    ws.row_dimensions[6].height = 18
    ws.row_dimensions[7].height = 28
    
    # ========== EXECUTIVE SUMMARY BOX ==========
    ws['B9'] = "EXECUTIVE SUMMARY"
    ws['B9'].font = header_font
    ws['B9'].fill = header_fill
    ws.merge_cells('B9:F9')
    ws['B9'].alignment = Alignment(horizontal='center', vertical='center')
    
    summary_text = (
        f"• {len(investigations[investigations['escalation_required'] == 'Yes'])} investigations escalated for senior review\n"
        f"• {len(customers[customers['risk_category'] == 'High'])} high-risk customers under enhanced monitoring\n"
        f"• {len(alerts[alerts['severity'] == 'High'])} high-severity alerts require immediate attention\n"
        f"• {len(customers[customers['kyc_status'] == 'Pending'])} KYC reviews pending remediation\n"
        f"• High-value transaction activity observed across {len(alerts[alerts['alert_type'] == 'High Value Transaction'])} cases"
    )
    
    ws.merge_cells('B10:F14')
    ws['B10'] = summary_text
    ws['B10'].font = Font(size=9, color="2C3E50")
    ws['B10'].alignment = Alignment(wrap_text=True, vertical='top')
    ws['B10'].fill = PatternFill(start_color="FDFEFE", end_color="FDFEFE", fill_type="solid")
    ws['B10'].border = thin_border
    
    # ========== CHART DATA (hidden) ==========
    ws_data = wb.create_sheet("_ChartData")
    
    # Alerts by Severity
    ws_data['A1'] = 'Severity'
    ws_data['B1'] = 'Count'
    for i, (sev, cnt) in enumerate(alerts['severity'].value_counts().items()):
        ws_data[f'A{i+2}'] = sev
        ws_data[f'B{i+2}'] = cnt
    
    # Alerts by Type
    ws_data['D1'] = 'Type'
    ws_data['E1'] = 'Count'
    for i, (typ, cnt) in enumerate(alerts['alert_type'].value_counts().items()):
        ws_data[f'D{i+2}'] = typ
        ws_data[f'E{i+2}'] = cnt
    
    # KYC Status
    ws_data['G1'] = 'KYC Status'
    ws_data['H1'] = 'Count'
    for i, (kyc, cnt) in enumerate(customers['kyc_status'].value_counts().items()):
        ws_data[f'G{i+2}'] = kyc
        ws_data[f'H{i+2}'] = cnt
    
    # Escalation
    ws_data['J1'] = 'Status'
    ws_data['K1'] = 'Count'
    ws_data['J2'] = 'Escalated'
    ws_data['K2'] = len(investigations[investigations['escalation_required'] == 'Yes'])
    ws_data['J3'] = 'Not Escalated'
    ws_data['K3'] = len(investigations[investigations['escalation_required'] == 'No'])
    
    # High-Risk Geography
    high_risk_countries = ['KP', 'IR', 'SY', 'SD', 'CU', 'VE', 'ZW', 'MM']
    hr_geo = customers[customers['country'].isin(high_risk_countries)]
    ws_data['M1'] = 'Country'
    ws_data['N1'] = 'Count'
    for i, (country, cnt) in enumerate(hr_geo['country'].value_counts().head(8).items()):
        ws_data[f'M{i+2}'] = country
        ws_data[f'N{i+2}'] = cnt
    
    # ========== CHARTS - COMPACT LAYOUT ==========
    
    # Row 1: Alerts by Severity (Doughnut) + Alerts by Type (Horizontal Bar)
    doughnut = DoughnutChart()
    doughnut.title = "Alerts by Severity"
    doughnut.style = 10
    labels = Reference(ws_data, min_col=1, min_row=2, max_row=3)
    data = Reference(ws_data, min_col=2, min_row=1, max_row=3)
    doughnut.add_data(data, titles_from_data=True)
    doughnut.set_categories(labels)
    doughnut.width = 11
    doughnut.height = 8
    ws.add_chart(doughnut, "B16")
    
    bar1 = BarChart()
    bar1.type = "bar"
    bar1.style = 10
    bar1.title = "Alerts by Type"
    labels2 = Reference(ws_data, min_col=4, min_row=2, max_row=4)
    data2 = Reference(ws_data, min_col=5, min_row=1, max_row=4)
    bar1.add_data(data2, titles_from_data=True)
    bar1.set_categories(labels2)
    bar1.width = 13
    bar1.height = 8
    ws.add_chart(bar1, "E16")
    
    # Row 2: KYC Status (Column) + Escalation Summary (Column)
    bar2 = BarChart()
    bar2.type = "col"
    bar2.style = 10
    bar2.title = "KYC Status Distribution"
    labels3 = Reference(ws_data, min_col=7, min_row=2, max_row=5)
    data3 = Reference(ws_data, min_col=8, min_row=1, max_row=5)
    bar2.add_data(data3, titles_from_data=True)
    bar2.set_categories(labels3)
    bar2.width = 11
    bar2.height = 8
    ws.add_chart(bar2, "B28")
    
    bar3 = BarChart()
    bar3.type = "col"
    bar3.style = 10
    bar3.title = "Escalation Summary"
    labels4 = Reference(ws_data, min_col=10, min_row=2, max_row=3)
    data4 = Reference(ws_data, min_col=11, min_row=1, max_row=3)
    bar3.add_data(data4, titles_from_data=True)
    bar3.set_categories(labels4)
    bar3.width = 13
    bar3.height = 8
    ws.add_chart(bar3, "E28")
    
    # Row 3: High-Risk Geography (Full width horizontal bar)
    bar4 = BarChart()
    bar4.type = "bar"
    bar4.style = 10
    bar4.title = "High-Risk Geography Activity"
    labels5 = Reference(ws_data, min_col=13, min_row=2, max_row=9)
    data5 = Reference(ws_data, min_col=14, min_row=1, max_row=9)
    bar4.add_data(data5, titles_from_data=True)
    bar4.set_categories(labels5)
    bar4.width = 24
    bar4.height = 8
    ws.add_chart(bar4, "B40")
    
    # ========== DATA SHEETS ==========
    
    # Sheet 1: Alert Queue
    ws1 = wb.create_sheet("Alert Queue")
    for r in dataframe_to_rows(alerts.head(100), index=False, header=True):
        ws1.append(r)
    ws1.auto_filter.ref = ws1.dimensions
    ws1.freeze_panes = "A2"
    
    # Sheet 2: Investigation Tracker
    ws2 = wb.create_sheet("Investigation Tracker")
    for r in dataframe_to_rows(investigations, index=False, header=True):
        ws2.append(r)
    ws2.auto_filter.ref = ws2.dimensions
    ws2.freeze_panes = "A2"
    
    # Sheet 3: High-Risk Customers
    ws3 = wb.create_sheet("High-Risk Customers")
    high_risk = customers[customers['risk_category'] == 'High']
    for r in dataframe_to_rows(high_risk, index=False, header=True):
        ws3.append(r)
    ws3.auto_filter.ref = ws3.dimensions
    ws3.freeze_panes = "A2"
    
    # Sheet 4: Escalation Summary
    ws4 = wb.create_sheet("Escalation Summary")
    summary_data = {
        'Metric': ['Total Escalated', 'EDD Required', 'High Risk', 'False Positives'],
        'Count': [
            len(investigations[investigations['escalation_required'] == 'Yes']),
            len(investigations[investigations['recommendation'] == 'Enhanced Due Diligence Required']),
            len(investigations[investigations['risk_level'] == 'High']),
            len(investigations[investigations['recommendation'] == 'False Positive'])
        ]
    }
    ws4.append(['Escalation Summary'])
    for r in dataframe_to_rows(pd.DataFrame(summary_data), index=False, header=True):
        ws4.append(r)
    
    # Sheet 5: KYC Review
    ws5 = wb.create_sheet("KYC Review")
    kyc_review = customers[['customer_id', 'kyc_status', 'onboarding_date', 'risk_category']].head(200)
    for r in dataframe_to_rows(kyc_review, index=False, header=True):
        ws5.append(r)
    ws5.auto_filter.ref = ws5.dimensions
    ws5.freeze_panes = "A2"
    
    # Sheet 6: Suspicious Transactions
    ws6 = wb.create_sheet("Suspicious Transactions")
    high_value = transactions[transactions['amount'] > 200000].head(200)
    susp_txn = high_value[['step', 'type', 'amount', 'nameOrig']]
    susp_txn.columns = ['day', 'transaction_type', 'amount', 'customer_id']
    for r in dataframe_to_rows(susp_txn, index=False, header=True):
        ws6.append(r)
    ws6.auto_filter.ref = ws6.dimensions
    ws6.freeze_panes = "A2"
    
    # Hide helper sheet
    ws_data.sheet_state = 'hidden'
    
    wb.save('reports/AML_Investigation_Workbook.xlsx')
    print("Excel workbook saved: reports/AML_Investigation_Workbook.xlsx")

def generate_word_report():
    alerts = pd.read_csv(ALERTS_PATH)
    investigations = pd.read_csv(INVESTIGATIONS_PATH)
    customers = pd.read_csv(CUSTOMERS_PATH)
    
    doc = Document()
    
    doc.add_heading('AML Investigation Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(f"""This report summarizes AML monitoring activities for the review period. A total of 
{len(alerts)} alerts were generated from {len(customers)} customers, with {len(investigations[investigations['escalation_required'] == 'Yes'])} cases requiring escalation. 
The monitoring system identified suspicious activity patterns requiring enhanced due diligence.""")
    
    doc.add_heading('2. AML Alert Overview', level=1)
    alert_summary = alerts.groupby('alert_type').size().reset_index(name='count')
    doc.add_paragraph("Alert breakdown by type:")
    for _, row in alert_summary.iterrows():
        doc.add_paragraph(f"  - {row['alert_type']}: {row['count']} alerts")
    
    doc.add_heading('3. Investigation Findings', level=1)
    findings = investigations['investigation_findings'].value_counts().head(5)
    for finding, count in findings.items():
        doc.add_paragraph(f"  - {finding}: {count} cases")
    
    doc.add_heading('4. Escalation Analysis', level=1)
    escalated = investigations[investigations['escalation_required'] == 'Yes']
    doc.add_paragraph(f"""{len(escalated)} cases were escalated for senior review. These cases typically involved 
high-risk customers, large transaction amounts, or suspicious activity patterns requiring additional scrutiny.""")
    
    doc.add_heading('5. High-Risk Customer Review', level=1)
    high_risk = customers[customers['risk_category'] == 'High']
    doc.add_paragraph(f"""{len(high_risk)} customers classified as high-risk require ongoing enhanced monitoring. 
These accounts are subject to additional transaction scrutiny and periodic compliance reviews.""")
    
    doc.add_heading('6. KYC Compliance Concerns', level=1)
    pending = customers[customers['kyc_status'] == 'Pending']
    rejected = customers[customers['kyc_status'] == 'Rejected']
    doc.add_paragraph(f"""KYC compliance gaps identified: {len(pending)} pending, {len(rejected)} rejected. 
These customers require immediate KYC remediation to maintain compliance standards.""")
    
    doc.add_heading('7. Recommendations', level=1)
    doc.add_paragraph("""1. Enhanced Due Diligence for PEP customers and high-risk geography activity\n2. Ongoing monitoring of escalated alerts\n3. KYC remediation for pending/rejected accounts\n4. Regular review of suspicious transaction patterns""")
    
    doc.save('reports/AML_Investigation_Report.docx')
    print("Word report saved: reports/AML_Investigation_Report.docx")

def main():
    print("=" * 60)
    print("AML DOCUMENTATION GENERATION")
    print("=" * 60)
    generate_excel_workbook()
    generate_word_report()

if __name__ == "__main__":
    main()